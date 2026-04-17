"""Text extraction from PDF, DOCX, TXT, and image files."""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional

from .config import SETTINGS

logger = logging.getLogger(__name__)

SUPPORTED_TEXT_EXTS = {".pdf", ".docx", ".doc", ".txt"}
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}
SUPPORTED_EXTS = SUPPORTED_TEXT_EXTS | SUPPORTED_IMAGE_EXTS


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            try:
                t = page.extract_text() or ""
            except Exception as e:  # noqa: BLE001
                logger.debug("pdfplumber page error in %s: %s", path, e)
                t = ""
            if t:
                chunks.append(t)
    return "\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as e:  # noqa: BLE001
        logger.debug("docx open error %s: %s", path, e)
        return ""
    paragraphs = [p.text for p in document.paragraphs if p.text]
    # Also pull text from tables (resumes often use tables for layout).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _extract_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


_OCR_UNAVAILABLE_REASON: Optional[str] = None


def ocr_available() -> tuple[bool, str]:
    """Return (is_available, reason). Caches the result on the module."""
    global _OCR_UNAVAILABLE_REASON
    if _OCR_UNAVAILABLE_REASON is not None:
        return (_OCR_UNAVAILABLE_REASON == ""), _OCR_UNAVAILABLE_REASON
    # An on-disk OCR cache (from scripts/ocr_images.py) is enough.
    try:
        from .ocr_cache import OCR_CACHE_DIR
        if OCR_CACHE_DIR.exists() and any(OCR_CACHE_DIR.iterdir()):
            _OCR_UNAVAILABLE_REASON = ""
            return True, ""
    except Exception:  # noqa: BLE001
        pass
    # Otherwise we need Tesseract on the system.
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except ImportError as e:
        _OCR_UNAVAILABLE_REASON = f"pytesseract/Pillow not importable: {e}"
        return False, _OCR_UNAVAILABLE_REASON
    import shutil
    tc = SETTINGS.tesseract_cmd
    if tc and os.path.exists(tc):
        _OCR_UNAVAILABLE_REASON = ""
        return True, ""
    if shutil.which("tesseract"):
        _OCR_UNAVAILABLE_REASON = ""
        return True, ""
    _OCR_UNAVAILABLE_REASON = (
        "no OCR cache found and tesseract binary not on PATH. "
        "Run `python -m scripts.ocr_images` first, or install Tesseract."
    )
    return False, _OCR_UNAVAILABLE_REASON


def _extract_image(path: Path) -> str:
    # 1) prefer the on-disk OCR cache produced by `scripts/ocr_images.py`.
    try:
        from .ocr_cache import read_cached
        cached = read_cached(path)
        if cached is not None:
            return cached.strip()
    except Exception:  # noqa: BLE001
        pass

    if not SETTINGS.enable_ocr:
        return ""
    ok, reason = ocr_available()
    if not ok:
        logger.debug("OCR unavailable (%s); skipping %s", reason, path)
        return ""
    import pytesseract
    from PIL import Image
    if SETTINGS.tesseract_cmd and os.path.exists(SETTINGS.tesseract_cmd):
        pytesseract.pytesseract.tesseract_cmd = SETTINGS.tesseract_cmd
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img).strip()
    except Exception as e:  # noqa: BLE001
        logger.debug("OCR failed for %s: %s", path, e)
        return ""


def extract_text(file_path: str | Path) -> Optional[str]:
    """Extract text from a single resume file. Returns None if unsupported."""
    path = Path(file_path)
    if not path.exists():
        return None
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext in (".docx", ".doc"):
            return _extract_docx(path)
        if ext == ".txt":
            return _extract_txt(path)
        if ext in SUPPORTED_IMAGE_EXTS:
            return _extract_image(path)
    except Exception as e:  # noqa: BLE001
        logger.warning("Failed to extract text from %s: %s", path, e)
        return ""
    return None


def is_supported(path: str | Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTS
